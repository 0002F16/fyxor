from pathlib import Path
import textwrap
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
ASSETS = OUT / "fyxor-architecture-assets"
DOCX = OUT / "Fyxor_Website_Architecture_Client_Documentation.docx"

GREEN = "047857"
EMERALD = "059669"
MINT = "D1FAE5"
SOFT = "F8FAFC"
LINE = "DADCE0"
INK = "0F172A"
MUTED = "475569"


def font(size=28, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill="#0F172A"):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=6, align="center")
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=fnt, fill=fill, spacing=6, align="center")


def arrow(draw, start, end, color="#059669", width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 14, ey - 8), (ex - sign * 14, ey + 8)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 8, ey - sign * 14), (ex + 8, ey - sign * 14)]
    draw.polygon(pts, fill=color)


def save_architecture_diagram():
    img = Image.new("RGB", (1800, 1050), "#F8FAFC")
    d = ImageDraw.Draw(img)
    title = font(52, True)
    h = font(30, True)
    body = font(24)
    small = font(20)
    d.text((70, 54), "Fyxor Architecture at a Glance", font=title, fill=f"#{INK}")
    d.text((72, 122), "How the browser, local device, server, database, and AI provider work together", font=body, fill=f"#{MUTED}")

    boxes = {
        "user": (80, 240, 380, 430, "Client / Candidate", "Uses Fyxor in Chrome"),
        "extension": (500, 200, 900, 470, "Chrome Extension", "Popup, full-page app,\nbackground worker,\nLinkedIn content script"),
        "local": (1040, 190, 1420, 460, "Local Device", "chrome.storage.local\nDownloads\nTemporary browser state"),
        "server": (500, 630, 900, 900, "Fyxor API Server", "Express API\nBetter Auth\nPDF/DOCX parsing\nAI orchestration\nCV export"),
        "db": (1040, 635, 1420, 885, "Postgres", "Auth tables\nuser_data\nusage_events"),
        "ai": (1490, 470, 1740, 720, "AI Providers", "Gemini API\nOpenAI API\nCodex CLI option")
    }
    for key, (x1, y1, x2, y2, label, desc) in boxes.items():
        rounded(d, (x1, y1, x2, y2), 28, "white", f"#{LINE}", 3)
        d.text((x1 + 32, y1 + 28), label, font=h, fill=f"#{GREEN}")
        d.multiline_text((x1 + 32, y1 + 84), desc, font=body, fill=f"#{INK}", spacing=8)

    arrow(d, (380, 335), (500, 335))
    arrow(d, (900, 305), (1040, 305))
    arrow(d, (700, 470), (700, 630))
    arrow(d, (900, 765), (1040, 765))
    arrow(d, (900, 710), (1490, 590))
    arrow(d, (1490, 625), (900, 825), "#94A3B8", 3)
    d.text((420, 300), "clicks,\nforms,\nfiles", font=small, fill=f"#{MUTED}", align="center")
    d.text((930, 245), "local-first\nstate", font=small, fill=f"#{MUTED}", align="center")
    d.text((720, 525), "authenticated\nAPI calls", font=small, fill=f"#{MUTED}", align="center")
    d.text((930, 720), "sync + usage", font=small, fill=f"#{MUTED}")
    d.text((1140, 520), "structured JSON + validated schemas", font=small, fill=f"#{MUTED}")
    img.save(ASSETS / "architecture_flow.png")


def save_ai_pipeline_diagram():
    img = Image.new("RGB", (1800, 920), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 56), "Resume Processing and AI Structuring", font=font(50, True), fill=f"#{INK}")
    d.text((72, 120), "The upload is converted to text first; AI only structures and rewrites within strict schemas.", font=font(24), fill=f"#{MUTED}")
    steps = [
        ("1", "User source", "PDF, DOCX,\nor pasted text"),
        ("2", "Text extraction", "pdf-parse for PDF\nmammoth for DOCX"),
        ("3", "AI extraction", "Prompt + schema\nNo invented facts"),
        ("4", "Validation", "Zod parsing\nSkill categories + IDs"),
        ("5", "Editable profile", "User reviews, edits,\nand saves base CV"),
    ]
    x = 70
    card_w = 260
    gap = 95
    for num, label, desc in steps:
        rounded(d, (x, 250, x + card_w, 570), 26, "#F8FAFC", f"#{LINE}", 3)
        d.ellipse((x + 28, 278, x + 82, 332), fill=f"#{GREEN}")
        center_text(d, (x + 28, 278, x + 82, 332), num, font(25, True), "white")
        d.text((x + 28, 360), label, font=font(30, True), fill=f"#{INK}")
        d.multiline_text((x + 28, 412), desc, font=font(22), fill=f"#{MUTED}", spacing=8)
        if x < 1320:
            arrow(d, (x + card_w, 410), (x + card_w + gap - 12, 410))
        x += card_w + gap
    rounded(d, (160, 690, 1640, 815), 24, "#ECFDF5", f"#{MINT}", 2)
    d.text((200, 720), "Important client takeaway", font=font(28, True), fill=f"#{GREEN}")
    takeaway = "The file itself is not saved as a file. The app keeps extracted source text and the structured profile\nso the user can verify everything before tailoring."
    d.multiline_text((200, 763), takeaway, font=font(24), fill=f"#{INK}", spacing=7)
    img.save(ASSETS / "ai_pipeline.png")


def save_interaction_diagram():
    img = Image.new("RGB", (1800, 1080), "#F8FAFC")
    d = ImageDraw.Draw(img)
    d.text((70, 56), "Main User Interaction Map", font=font(50, True), fill=f"#{INK}")
    rows = [
        ("Account", "Sign up, sign in, sign out, sync user data"),
        ("Onboarding", "Upload PDF/DOCX, paste text, structure with AI, or enter manually"),
        ("Base Resume", "Inline edit profile, reorder sections, review resume-strength tips"),
        ("Job Capture", "LinkedIn auto-detect, page dialog, popup, or right-click selected text"),
        ("Tailoring", "Generate tailored CV, monitor progress, open result"),
        ("Editing", "Inline edits, regenerate summary/experience/skills, handle unsupported-claim warnings"),
        ("Tracker", "Preview, open, duplicate, or delete application records"),
        ("Export", "Download tailored CV as PDF or DOCX"),
        ("Settings", "Provider, engine, server URL, output language, health check"),
    ]
    y = 160
    for i, (label, desc) in enumerate(rows):
        x = 105 if i % 2 == 0 else 965
        if i % 2 == 0 and i > 0:
            y += 190
        rounded(d, (x, y, x + 730, y + 140), 22, "white", f"#{LINE}", 2)
        d.text((x + 28, y + 24), label, font=font(29, True), fill=f"#{GREEN}")
        d.text((x + 28, y + 72), desc, font=font(23), fill=f"#{INK}")
    img.save(ASSETS / "interaction_map.png")


def save_local_server_split():
    img = Image.new("RGB", (1800, 950), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 55), "What Runs Locally vs. What Runs on the Server", font=font(48, True), fill=f"#{INK}")
    columns = [
        ("Local Chrome Extension", [
            "Renders full-page app, popup, and LinkedIn overlay",
            "Stores profile, drafts, applications, pending jobs, and auth token locally",
            "Reads LinkedIn job pages through a content script",
            "Adds right-click “Send selection to Fyxor” action",
            "Downloads generated PDF/DOCX files to the user’s device",
        ]),
        ("Fyxor API Server", [
            "Owns authentication and bearer-session validation",
            "Extracts text from PDF/DOCX uploads",
            "Calls selected AI provider or CCC engine",
            "Validates AI output against strict schemas",
            "Syncs profile/drafts/applications to Postgres",
            "Tracks monthly usage and exports CV files",
        ])
    ]
    for idx, (title, items) in enumerate(columns):
        x = 90 + idx * 860
        rounded(d, (x, 180, x + 760, 815), 30, "#F8FAFC" if idx == 0 else "#ECFDF5", f"#{LINE}", 3)
        d.text((x + 36, 225), title, font=font(34, True), fill=f"#{GREEN}")
        yy = 300
        for item in items:
            d.ellipse((x + 40, yy + 8, x + 54, yy + 22), fill=f"#{EMERALD}")
            wrapped = "\n".join(textwrap.wrap(item, width=54))
            d.multiline_text((x + 76, yy), wrapped, font=font(22), fill=f"#{INK}", spacing=5)
            yy += 82
    img.save(ASSETS / "local_server_split.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Fyxor Table"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_figure(doc, image_name, caption, width=6.2):
    doc.add_picture(str(ASSETS / image_name), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(GREEN)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)

    table_style = styles.add_style("Fyxor Table", 3)
    table_style.font.name = "Calibri"
    table_style.font.size = Pt(9)


def header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Fyxor client documentation"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = "Website architecture, APIs, and interaction inventory"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    save_architecture_diagram()
    save_ai_pipeline_diagram()
    save_interaction_diagram()
    save_local_server_split()

    doc = Document()
    style_document(doc)
    header_footer(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Fyxor Website & Extension Architecture")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Client-facing documentation: what runs locally, what runs on the server, APIs used, and every major user interaction")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for client review | Current implementation snapshot")
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_figure(doc, "architecture_flow.png", "Figure 1. High-level architecture across the browser, local storage, server, database, and AI provider.", 6.3)
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    paragraph(doc, "Fyxor is a Chrome extension plus companion API server for creating a verified base CV, reading job descriptions, generating tailored CVs, editing those CVs, tracking applications, and exporting the result as PDF or DOCX.")
    paragraph(doc, "The product is local-first: the extension keeps the working profile, drafts, applications, pending jobs, and session token on the user’s device. The server handles the sensitive and heavier work: authentication, cloud sync, file parsing, AI calls, usage tracking, and document export.")
    paragraph(doc, "For a non-technical stakeholder, the most important mental model is simple: the browser is the workspace, the API server is the processing hub, Postgres is the account-backed backup, and AI is called only through the server-side orchestration layer.")

    doc.add_heading("What the Client Can See Without Opening the Extension", level=2)
    add_figure(doc, "home.png", "Figure 2. Home screen after a tailored CV exists: the user can reopen recent work quickly.", 6.2)
    add_figure(doc, "onboarding.png", "Figure 3. Onboarding starts from a real CV file or pasted text, then structures the profile with AI.", 6.2)
    doc.add_page_break()
    add_figure(doc, "resume.png", "Figure 4. Base resume editor: the verified source CV that future tailoring is based on.", 6.2)
    add_figure(doc, "tracker.png", "Figure 5. Application tracker with resume preview, open, duplicate, and delete actions.", 6.2)
    doc.add_page_break()
    add_figure(doc, "editor.png", "Figure 6. Tailored CV editor: inline edits, regeneration, resume strength, and PDF/DOCX export.", 6.2)
    add_figure(doc, "account.png", "Figure 7. Account and settings: profile summary, sync state, language, provider, engine, and server configuration.", 6.2)
    doc.add_page_break()

    doc.add_heading("2. Architecture Overview", level=1)
    add_figure(doc, "local_server_split.png", "Figure 8. Responsibility split between the Chrome extension and the API server.", 6.3)
    doc.add_heading("Local Chrome Extension", level=2)
    add_bullets(doc, [
        "Renders the full-page app: Home, onboarding, base resume, applications tracker, tailored CV editor, and account settings.",
        "Renders the Chrome toolbar popup used to import a detected or selected job and start tailoring.",
        "Runs a background service worker for context-menu handling, tailoring progress, badge updates, and opening the full-page app.",
        "Runs a LinkedIn content script on LinkedIn job pages to detect job details and display a small page overlay/dialog.",
        "Stores local working state using chrome.storage.local, with a localStorage fallback when opened in a normal browser.",
        "Downloads exported CV files through Chrome downloads or a browser download link."
    ])

    doc.add_heading("Server-Side API", level=2)
    add_bullets(doc, [
        "Runs Express with CORS restricted to Chrome extension origins, localhost, and 127.0.0.1.",
        "Runs Better Auth with bearer-token sessions for Chrome extension compatibility.",
        "Parses uploaded PDF and DOCX files into text.",
        "Calls AI providers or the optional CCC engine to structure, tailor, and regenerate CV content.",
        "Validates all AI output using shared Zod schemas before returning it to the extension.",
        "Stores account-backed user data and usage events in Postgres.",
        "Exports final CVs to PDF or DOCX."
    ])

    doc.add_heading("External Services and Libraries", level=2)
    add_table(doc, ["Area", "Technology / API", "Used For"], [
        ("Authentication", "Better Auth + bearer plugin", "Email/password sign-up, sign-in, sign-out, and bearer token validation."),
        ("Database", "Postgres via pg Pool", "Self-hosted auth tables, user_data, and usage_events."),
        ("AI providers", "Gemini API, OpenAI API, Codex CLI provider code", "Resume extraction, tailoring, and regeneration. Current stored defaults favor Gemini."),
        ("File parsing", "pdf-parse, mammoth", "Extract text from PDF and DOCX uploads before AI processing."),
        ("Document export", "pdfkit, docx", "Generate downloadable PDF and DOCX CV files."),
        ("Browser integration", "Chrome MV3 APIs", "storage, tabs, downloads, contextMenus, scripting, activeTab, action popup."),
        ("LinkedIn integration", "Content script DOM parsing", "Read job title, company, location, description, BPR blocks, and JSON-LD when present."),
    ], [1.15, 1.85, 3.2])

    doc.add_heading("3. Data and Storage Model", level=1)
    paragraph(doc, "Fyxor keeps a complete local working copy and syncs account-level data to the server. Settings remain device-local so a user can point one device at a local server and another at a hosted server.")
    add_table(doc, ["Data", "Stored Locally", "Synced to Server", "Notes"], [
        ("Auth session", "Yes", "Session is validated server-side", "Token is stored locally and sent as Authorization: Bearer <token>."),
        ("Base profile", "Yes", "Yes", "The source profile used for all future tailoring."),
        ("Draft tailored CVs", "Yes", "Yes", "Saved in drafts by tailored CV ID."),
        ("Application tracker", "Yes", "Yes", "Each record includes the job and tailored CV."),
        ("Pending job", "Yes", "No", "Used while moving from LinkedIn/selection capture to the popup/editor."),
        ("Settings", "Yes", "No", "API base URL, provider, tailoring engine, UI flags, output language reference."),
        ("Usage events", "No", "Yes", "Server records actions like extract, tailor, regenerate."),
        ("Uploaded file bytes", "Temporary request only", "No file storage in app tables", "The server extracts text and returns text/profile data."),
    ], [1.25, 1.05, 1.15, 3.0])

    doc.add_heading("Postgres Tables Owned by the App", level=2)
    add_table(doc, ["Table", "Purpose", "Main Fields"], [
        ("user_data", "Cloud sync for the user’s profile and work.", "user_id, profile jsonb, drafts jsonb, applications jsonb, updated_at"),
        ("usage_events", "Monthly usage accounting and future free-tier limits.", "id, user_id, action, created_at, meta jsonb"),
    ], [1.3, 2.4, 2.8])
    paragraph(doc, "Better Auth also creates its own user, session, account, and verification tables through its migration command.")

    doc.add_heading("4. Resume Processing and AI", level=1)
    add_figure(doc, "ai_pipeline.png", "Figure 9. Resume import and structuring pipeline.", 6.3)
    doc.add_heading("Onboarding Extraction", level=2)
    add_numbered(doc, [
        "The user uploads a PDF/DOCX or pastes raw CV text.",
        "If a file is uploaded, the extension sends base64 content and filename to /api/profile/parse-file.",
        "The server extracts plain text with pdf-parse for PDF or mammoth for DOCX.",
        "The user clicks Structure with AI; the extension sends text and output language to /api/profile/extract.",
        "The server calls the configured generator with an extraction prompt and strict base-profile schema.",
        "The AI returns JSON, not prose. The API parses and validates the JSON, folds skill categories into the app’s internal format, stores rawText, and returns an editable profile.",
        "The user reviews and edits the profile before completing onboarding."
    ])
    doc.add_heading("AI Guardrails", level=2)
    add_bullets(doc, [
        "Do not invent employers, roles, dates, tools, skills, qualifications, achievements, metrics, or responsibilities.",
        "Leave unknown education fields empty instead of guessing.",
        "Keep skills grounded in the source text.",
        "For tailored CVs, preserve sourceExperienceId and sourceBulletIndexes so rewritten bullets trace back to source evidence.",
        "Return unsupportedClaims when proposed wording may overreach the source evidence."
    ])
    doc.add_heading("Tailoring and Regeneration", level=2)
    paragraph(doc, "Tailoring takes the verified base profile plus a job description and produces a tailored CV. If the CCC engine is configured and selected, the server runs that multi-step engine. If not, it falls back to the built-in single-pass generator. Regeneration is narrower: it updates only summary, experience, or skills while preserving the rest of the CV.")

    doc.add_heading("5. API Inventory", level=1)
    add_table(doc, ["Endpoint", "Auth", "Client Action", "What It Does"], [
        ("GET /health", "No", "Account settings: Test selected provider", "Reports API health, provider configuration, model, and CCC engine availability."),
        ("POST /api/auth/sign-up/email", "No", "Create account", "Creates email/password account through Better Auth and returns bearer token."),
        ("POST /api/auth/sign-in/email", "No", "Sign in", "Authenticates account and returns bearer token."),
        ("POST /api/auth/sign-out", "Bearer", "Sign out", "Best-effort server sign-out; local state is cleared."),
        ("POST /api/profile/parse-file", "No in current route", "Upload PDF/DOCX", "Extracts raw text from uploaded file bytes."),
        ("POST /api/profile/extract", "Bearer", "Structure with AI", "Turns raw CV text into a structured base profile."),
        ("POST /api/cvs/tailor", "Bearer", "Generate tailored CV", "Creates a tailored CV using CCC engine or built-in AI provider."),
        ("POST /api/cvs/regenerate", "Bearer", "Regenerate section", "Regenerates summary, experience, or skills only."),
        ("GET /api/data/sync", "Bearer", "Sign-in data pull", "Returns profile, drafts, and applications for this user."),
        ("PUT /api/data/sync", "Bearer", "Debounced sync", "Upserts local profile, drafts, and applications into Postgres."),
        ("GET /api/data/usage", "Bearer", "Account usage display", "Returns current-month usage counts by action."),
        ("POST /api/cvs/export?format=pdf|docx", "No bearer check in current route", "Download PDF/DOCX", "Generates a CV export and returns a downloadable file."),
    ], [1.55, 0.7, 1.55, 2.7])

    doc.add_heading("6. Complete Interaction Inventory", level=1)
    add_figure(doc, "interaction_map.png", "Figure 10. Main surfaces and actions available to users.", 6.3)
    doc.add_heading("Account and Entry", level=2)
    add_bullets(doc, [
        "Create account with name, email, and password.",
        "Sign in with email and password.",
        "On sign-in, pull cloud profile/drafts/applications and adopt them locally.",
        "Sign out, clearing local profile, drafts, applications, pending jobs, and session token.",
        "Unauthenticated users see the sign-in/sign-up gate before any onboarding or tailoring screen."
    ])
    doc.add_heading("Onboarding", level=2)
    add_bullets(doc, [
        "Choose upload flow or manual resume creator from the welcome screen.",
        "Upload PDF or DOCX.",
        "Paste raw CV text.",
        "Run Structure with AI after enough text is present.",
        "Skip AI and enter details manually.",
        "Edit basics: name, email, phone, location, LinkedIn, target role, summary.",
        "Add, edit, or remove experience entries and bullets.",
        "Group skills by category; add, rename, edit, or remove categories.",
        "Add, edit, or remove education entries including school, degree, location, graduation date, GPA, honors, coursework.",
        "Add certifications and languages.",
        "Finish setup and move to the pin-Fyxor instruction screen."
    ])
    doc.add_heading("Base Resume", level=2)
    add_bullets(doc, [
        "View the base CV as a resume canvas.",
        "Click lines to edit text directly.",
        "Edit the headline / target role.",
        "Use resume-strength tips and dismiss individual checks.",
        "Redo guided setup from the resume screen.",
        "Autosave base resume edits after a short delay."
    ])
    doc.add_heading("Job Capture", level=2)
    add_bullets(doc, [
        "On LinkedIn job pages, the content script scans DOM, BPR blocks, JSON-LD, and visible page text.",
        "A small LinkedIn indicator shows whether a job is detected, scanning, or unsupported.",
        "The LinkedIn page dialog can generate a tailored CV from the detected job or prompt onboarding if no base profile exists.",
        "The toolbar popup reads the current LinkedIn job when possible.",
        "On any website, the user can highlight job text, right-click, and choose Send selection to Fyxor.",
        "If the popup cannot open, the extension badge shows a pending-job indicator."
    ])
    doc.add_heading("Tailoring Workflow", level=2)
    add_bullets(doc, [
        "Open popup with imported or detected job.",
        "Review job title, company, and description preview.",
        "Generate tailored CV.",
        "Progress state is stored in local extension state so the popup can close while the background worker continues.",
        "When tailoring succeeds, a draft and application record are created.",
        "Open the tailored CV editor from the popup or tracker."
    ])
    doc.add_heading("Tailored CV Editor", level=2)
    add_bullets(doc, [
        "Inline edit any line on the resume canvas.",
        "Regenerate summary, a specific experience, or skills.",
        "Reorder sections from the resume canvas controls.",
        "Review unsupported-claim warnings when AI flags wording that may be too strong.",
        "See resume-strength guidance.",
        "Accept an education-first nudge for candidates with thin work history.",
        "Export the tailored CV as PDF or DOCX.",
        "Autosave edits back to drafts and the linked application record."
    ])
    doc.add_heading("Applications Tracker", level=2)
    add_bullets(doc, [
        "List all tailored application records.",
        "Select an application to preview its tailored resume.",
        "Open a tailored CV for editing.",
        "Duplicate an application and its CV draft.",
        "Delete an application after confirmation."
    ])
    doc.add_heading("Account Settings", level=2)
    add_bullets(doc, [
        "Edit base resume or restart guided setup.",
        "Choose output language: English or Polish.",
        "View sync state and monthly usage when available.",
        "Choose AI provider: Gemini API, OpenAI API, or Codex local option in provider code.",
        "Choose tailoring engine: CCC studio engine when detected, or built-in single pass.",
        "Change API server URL.",
        "Save settings and test selected provider."
    ])

    doc.add_heading("7. Security, Privacy, and Operational Notes", level=1)
    add_bullets(doc, [
        "The OpenAI/Gemini keys stay on the server, not inside the extension.",
        "The extension uses bearer tokens because MV3 extensions are unreliable with cross-origin cookies.",
        "The manifest currently grants storage, tabs, downloads, contextMenus, scripting, activeTab, LinkedIn host access, localhost API access, and the configured hosted API origin.",
        "CORS allows chrome-extension origins plus localhost/127.0.0.1 development origins.",
        "The JSON body limit is 15 MB, which is relevant for uploaded resume files encoded as base64.",
        "Cloud sync is debounced and last-write-wins. If offline, local state remains usable and sync resumes on later changes.",
        "The parse-file and export routes are not bearer-gated in the current implementation. They do not directly expose stored user data, but adding auth would make the boundary stricter.",
        "The README mentions Codex local as a provider, while current shared-state migration moves old codex-local settings to Gemini. This is worth aligning in product documentation before external handoff."
    ])

    doc.add_heading("8. Client-Friendly Summary", level=1)
    paragraph(doc, "Fyxor is already more than a simple form. It is a browser-based CV workspace with a server-side AI and export layer. The local extension owns the user experience and day-to-day working state. The server owns trusted processing: login, file parsing, AI calls, schema validation, synchronization, usage counts, and final exports.")
    paragraph(doc, "The important business strength is that AI is constrained by a verified base profile and strict schemas. The user is always shown editable output before it becomes the source for tailoring, and the final editor keeps the client in control with inline edits, regeneration, warnings, and exports.")

    doc.save(DOCX)


if __name__ == "__main__":
    build_doc()
    print(DOCX)
